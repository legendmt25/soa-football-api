from datetime import date
from io import BytesIO
from src.repositories import BettingRepository
from src.integrations import FootballClient

from src.models import BetCreate
from src.schemas import Bet, BetMatch

import docx

class BettingService:
    def __init__(self, bettingRepository: BettingRepository, footballClient: FootballClient):
        self.footballClient = footballClient
        self.bettingRepository = bettingRepository

    def findAll(self):
        return self.bettingRepository.findAll()

    def findById(self, id: int):
        return self.bettingRepository.findById(id)

    def placeBet(self, bet: BetCreate, userId: str):
        return self.bettingRepository.create(
            Bet(
                userId = userId,
                createdAt = date.today(),
                matches = [BetMatch(playType=i.playType, matchId=i.matchId) for i in bet.matches],
                price = bet.price
            )
        )
    
    def printTicket(self, id: int, jwttoken: str):
        stream = BytesIO()
        document = docx.Document()
        bet = self.findById(id)
        matchDict = dict()

        document.add_heading('Bet ticket' + ' ' + str(bet.id) + ': ', 0)
        document.add_paragraph('User: ' + str(bet.userId))
        for match in bet.matches:
            if not match.matchId in matchDict:
                matchDict[match.matchId] = self.footballClient.getMatchById(match.matchId, jwttoken)
            temp = matchDict.get(match.matchId)
            print(temp)
            document.add_paragraph(temp['home_team']['name'] + '-' + temp['away_team']['name'] + ' | ' + temp['match_start'] + ' | ' + match.playType)

        document.add_paragraph('Win: ' + str(100 * bet.price) + '$')

        document.save(stream)
        stream.flush()
        stream.seek(0)
        return stream

